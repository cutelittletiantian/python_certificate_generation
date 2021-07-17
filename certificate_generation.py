import os
import pandas as pd
from docx import shared
from pandas import DataFrame
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# 配置基本参数
# 证件照资源路径
photoPath = r"resources\证件照"
# 生成证书保存路径
generationPath = r"resources\生成_证书文件_word"
if not os.path.exists(generationPath):
    os.mkdir(path=generationPath)
# 模板证书所在路径
templatePath = r"resources\证书模板.docx"
# 数据表所在路径
memberBookPath = r"resources\组织成员工作评价表.xlsx"

# 加载基本数据
# 所有照片的完整路径
photoItemPathList = []
for parentDir, subDirList, photoList in os.walk(photoPath):
    for photoItem in photoList:
        photoItemPathList.append(os.path.join(parentDir, photoItem))
# 读取excel表格，认定信息表
dfMember = pd.read_excel(io=memberBookPath, sheet_name=0,
                         usecols=["姓名", "部门", "职务", "认定级别\n（优秀、良好、合格、不合格）"])  # type: DataFrame

# 逐行遍历dfMember表
for memberIndex, seriesMember in dfMember.iterrows():
    # 取出姓名、部门名称、职务、评定等级
    name, department, career, result = seriesMember.values.tolist()
    # 不合格不予生成证书
    if result == "不合格":
        continue

    # 打开模板文件，为其生成证书
    docCertTemplate = docx.Document(docx=templatePath)  # type: Document

    # 替换状态标签
    flagDollarDetected = False

    # 扫描整个文档，替换其中的$xxx$标定点
    for paraIndex, para in enumerate(docCertTemplate.paragraphs):
        for runIndex, run in enumerate(para.runs):
            # 第一次遇到$号
            if run.text == "$" and not flagDollarDetected:
                run.text = run.text.replace("$", "")
                flagDollarDetected = True
            # 夹在$符号中间，执行替换
            elif run.text != "$" and flagDollarDetected:
                run.text = run.text.replace("姓名", name)
                run.text = run.text.replace("部门名称", department)
                run.text = run.text.replace("职务", career)
                run.text = run.text.replace("评定等级", result)
            # 第二次遇到$号闭合
            elif run.text == "$" and flagDollarDetected:
                run.text = run.text.replace("$", "")
                flagDollarDetected = False

    # 进入一个隐藏表格，替换其中的$照片$标定点为1寸照片(2.5*3.5cm)
    for paraIndex, para in enumerate(docCertTemplate.tables[0].cell(0, 0).paragraphs):
        for runIndex, run in enumerate(para.runs):
            # 第一次遇到$号
            if run.text == "$" and not flagDollarDetected:
                run.text = run.text.replace("$", "")
                flagDollarDetected = True
            # 夹在$符号中间，执行替换
            elif run.text == "照片" and flagDollarDetected:
                run.text = run.text.replace("照片", "")

                for item in photoItemPathList:
                    # 找到
                    if name in item:
                        # 打开图片stream
                        with open(file=item, mode="rb") as imgBuffer:
                            # 添加照片，调整为1寸照
                            run.add_picture(image_path_or_stream=imgBuffer,
                                            width=docx.shared.Cm(2.5))
                        break
                # end of for: 一直没找到
                else:
                    run.text = "404 NOT FOUND"
            # 第二次遇到$号闭合
            elif run.text == "$" and flagDollarDetected:
                run.text = run.text.replace("$", "")
                flagDollarDetected = False

    # 生成的证书保存完毕
    docCertPath = os.path.join(generationPath, f"{name}_{department}{career}_{result}证书.docx")
    # 点击“另存”，以【{姓名}_{部门}{职务}_{评定等级}证书.docx】为格式命名，放在生成目录下
    docCertTemplate.save(path_or_stream=docCertPath)

    print(f"生成完毕 {name} 同学的证书~~~，保存到路径'{docCertPath}'中。")
